import streamlit as st
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import io
import docx.oxml
import docx.oxml.ns

# --- Set page config ---
st.set_page_config(page_title="📸 MED Pictures Generator", layout="wide")
st.title(u"\U0001F4F8 MED PICTURES Word Document Generator")  # 📸

# --- Crop size mapping (in inches based on layout + 2cm margin logic) ---
crop_options = {
    # PORTRAIT A4 (6.69 x 10.1 inches)
    "Portrait: 2 x 1": (4.95, 6.69),
    "Portrait: 2 x 2": (4.95, 3.25),
    "Portrait: 2 x 3": (4.95, 2.13),
    "Portrait: 3 x 1": (3.27, 6.69),
    "Portrait: 3 x 2": (3.27, 3.25),
    "Portrait: 3 x 3": (3.27, 2.13),
    # LANDSCAPE A4 (10.1 x 6.69 inches)
    "Landscape: 2 x 1": (3.25, 10.1),
    "Landscape: 2 x 2": (3.25, 4.95),
    "Landscape: 2 x 3": (3.25, 3.27),
    "Landscape: 3 x 1": (2.13, 10.1),
    "Landscape: 3 x 2": (2.13, 4.95),
    "Landscape: 3 x 3": (2.13, 3.27),
}

# --- User input form ---
with st.form("image_form"):
    st.subheader("🛠️ Document Configuration")
    title = st.text_input("Project Title")
    contractor = st.text_input("Contractor Name")
    orientation = st.selectbox("Page Orientation", ["Portrait", "Landscape"])

    uploaded_images = st.file_uploader("Upload Images", type=["jpg", "jpeg", "png"], accept_multiple_files=True)

    crop_selections = []
    if uploaded_images:
        st.markdown("### ✂️ Select Crop Size per Image")
        for idx, image_file in enumerate(uploaded_images):
            col1, col2 = st.columns([1, 2])
            with col1:
                st.image(image_file, caption=f"Image {idx+1}", use_container_width=True)
            with col2:
                crop = st.selectbox(
                    f"Crop size for Image {idx+1}",
                    options=[k for k in crop_options if k.startswith(orientation)],
                    key=f"crop_{idx}"
                )
                crop_selections.append(crop)

    layout = st.selectbox("Grid Layout per Page", ["1 x 2", "1 x 3", "2 x 2", "2 x 3", "3 x 2", "3 x 3"])

    submitted = st.form_submit_button("📄 Generate Word Document")

# --- Document generator ---
def generate_doc(title, contractor, images, crop_sizes, layout, orientation):
    layout_map = {
        "2 x 1": (2, 1), "2 x 2": (2, 2), "2 x 3": (2, 3),
        "3 x 1": (3, 1), "3 x 2": (3, 2), "3 x 3": (3, 3)
    }
    rows, cols = layout_map[layout]
    images_per_page = rows * cols

    doc = Document()
    section = doc.sections[0]

    if orientation == 'Landscape':
        section.orientation = 1
        section.page_width, section.page_height = section.page_height, section.page_width

    section.top_margin = Inches(0.69)
    section.bottom_margin = Inches(0.69)
    section.left_margin = Inches(0.79)
    section.right_margin = Inches(0.79)

    usable_width = section.page_width - section.left_margin - section.right_margin
    gap_between_cols = Inches(0.197)  # 1 cm
    total_gap_width = gap_between_cols * (cols - 1)
    usable_table_width = usable_width - total_gap_width
    col_width = usable_table_width / cols

    for i in range(0, len(images), images_per_page):
        p = doc.add_paragraph()
        run1 = p.add_run("MED PICTURES: ")
        run1.font.color.rgb = RGBColor(255, 0, 0)
        run1.bold = True
        run2 = p.add_run(f"{title} by {contractor}")
        run2.bold = True

        table = doc.add_table(rows=rows, cols=cols)
        table.autofit = False

        for row in table.rows:
            for j, cell in enumerate(row.cells):
                cell.width = col_width
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                tc_pr = cell._element.get_or_add_tcPr()
                cell_margin_xml = docx.oxml.parse_xml(
                    '<w:tcMar '
                    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    '<w:top w:w="0"/><w:left w:w="0"/><w:bottom w:w="0"/><w:right w:w="0"/>'
                    '</w:tcMar>'
                )
                tc_pr.append(cell_margin_xml)

        for idx, image_file in enumerate(images[i:i + images_per_page]):
            r, c = divmod(idx, cols)
            cell = table.rows[r].cells[c]

            img = Image.open(image_file).convert("RGB")
            crop_label = crop_sizes[i + idx]
            width_in, height_in = crop_options[crop_label]
            target_px = (int(width_in * 300), int(height_in * 300))
            img = img.resize(target_px)

            img_stream = io.BytesIO()
            img.save(img_stream, format='PNG')
            img_stream.seek(0)

            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(img_stream, width=Inches(width_in))

        if i + images_per_page < len(images):
            doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Document output ---
if submitted:
    if not title or not contractor or not uploaded_images:
        st.error("Please provide all inputs.")
    else:
        word_doc = generate_doc(title, contractor, uploaded_images, crop_selections, layout, orientation)
        st.success("✅ Document ready!")
        st.download_button("📥 Download Word Document", word_doc, "MED_PICTURES.docx")
