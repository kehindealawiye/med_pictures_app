# === IMPORT LIBRARIES ===
import streamlit as st
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import io

# ✅ Set page config before any Streamlit output
st.set_page_config(page_title="MED Pictures Generator", layout="wide")

# === TITLE ===
st.title("\ud83d\udcf8 MED PICTURES Word Document Generator")

# === LAYOUT OPTIONS MAPPING ===
layout_options = {
    "1 x 2": (1, 2),
    "1 x 3": (1, 3),
    "2 x 2": (2, 2),
    "2 x 3": (2, 3),
    "3 x 2": (3, 2),
    "3 x 3": (3, 3),
}

# === DOCUMENT GENERATION FUNCTION ===
def generate_doc(title, contractor, images, layout, orientation, crop_sizes):
    rows, cols = layout_options[layout]
    images_per_page = rows * cols

    # Create Word document
    doc = Document()
    section = doc.sections[0]

    # Set orientation
    if orientation == 'Landscape':
        section.orientation = 1
        section.page_width, section.page_height = section.page_height, section.page_width

    # Set minimal page margins
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Calculate usable width of the page
    usable_width = section.page_width - section.left_margin - section.right_margin
    col_width = usable_width / cols

    # Map crop sizes to inches (W, H)
    crop_dimensions = {
        "11.03 x 9 (Portrait Short)": (Inches(3.54), Inches(4.34)),
        "13.26 x 9.3 (Portrait Tall)": (Inches(3.66), Inches(5.22)),
        "8.56 x 10.58 (Landscape)": (Inches(4.17), Inches(3.37)),
        "8.56 x 18.94 (Wide Landscape)": (Inches(7.46), Inches(3.37)),
    }

    # Process images and paginate
    for i in range(0, len(images), images_per_page):
        p = doc.add_paragraph()
        run1 = p.add_run("MED PICTURES: ")
        run1.font.color.rgb = RGBColor(255, 0, 0)
        run1.bold = True
        run2 = p.add_run(f"{title} by {contractor}")
        run2.bold = True

        table = doc.add_table(rows=rows, cols=cols)
        table.autofit = False

        for row in table.rows:
            for cell in row.cells:
                cell.width = col_width
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for idx, image_file in enumerate(images[i:i + images_per_page]):
            r, c = divmod(idx, cols)
            cell = table.rows[r].cells[c]

            crop_label = crop_sizes.get(image_file.name, "11.03 x 9 (Portrait Short)")
            target_width, target_height = crop_dimensions[crop_label]

            img = Image.open(image_file)
            img_ratio = img.width / img.height
            target_ratio = target_width / target_height

            # Auto-center crop to match aspect ratio
            if img_ratio > target_ratio:
                new_width = int(img.height * target_ratio)
                offset = (img.width - new_width) // 2
                img = img.crop((offset, 0, offset + new_width, img.height))
            else:
                new_height = int(img.width / target_ratio)
                offset = (img.height - new_height) // 2
                img = img.crop((0, offset, img.width, offset + new_height))

            image_stream = io.BytesIO()
            img.save(image_stream, format='PNG')
            image_stream.seek(0)

            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(image_stream, width=target_width)

        if i + images_per_page < len(images):
            doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# === USER INPUT FORM ===
with st.form("image_form"):
    st.subheader("\ud83d\udd27 Document Configuration")

    title = st.text_input("Project Title")
    contractor = st.text_input("Contractor Name")
    orientation = st.selectbox("Page Orientation", ["Portrait", "Landscape"])
    layout = st.selectbox("Image Layout", list(layout_options.keys()))
    uploaded_images = st.file_uploader("Upload Images", type=["jpg", "jpeg", "png"], accept_multiple_files=True)

    crop_sizes = {}
    if uploaded_images:
        st.markdown("### \u2702\ufe0f Select Crop Size per Image")
        for i, image_file in enumerate(uploaded_images):
            with st.container():
                cols = st.columns([1, 2])
                with cols[0]:
                    st.image(image_file, caption=f"Image {i+1}", use_column_width=True)
                with cols[1]:
                    crop = st.selectbox(
                        f"Crop size for Image {i+1}",
                        options=[
                            "11.03 x 9 (Portrait Short)",
                            "13.26 x 9.3 (Portrait Tall)",
                            "8.56 x 10.58 (Landscape)",
                            "8.56 x 18.94 (Wide Landscape)"
                        ],
                        key=f"crop_{i}"
                    )
                    crop_sizes[image_file.name] = crop

    submitted = st.form_submit_button("Generate Word Document")

# === DOCUMENT GENERATION AND DOWNLOAD ===
if submitted:
    if not title or not contractor or not uploaded_images:
        st.error("Please provide all required inputs.")
    else:
        word_file = generate_doc(title, contractor, uploaded_images, layout, orientation, crop_sizes)
        st.success("\u2705 Document ready!")
        st.download_button(
            label="\ud83d\udcc5 Download Document",
            data=word_file,
            file_name="MED_Pictures.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )